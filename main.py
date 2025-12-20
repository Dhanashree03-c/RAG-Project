from fastapi import FastAPI, UploadFile, Body
import uuid, os
from PIL import Image
import pytesseract
import pdfplumber
import docx
import requests
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from pymongo import MongoClient
from datetime import datetime, timezone
from googlesearch import search
from bs4 import BeautifulSoup
from readability import Document
from pydantic import BaseModel

#MONGODB ATLAS SETUP
MONGO_URI = "mongodb+srv://tankardhanashree05_db_user:qtlBmgKu7m9bppWd@cluster0.zdc8mpn.mongodb.net/?appName=Cluster0"
client = MongoClient(MONGO_URI)
db = client["document_db"]
collection = db["documents"]                    # document chunks + metadata
metrics_collection = db["metrics"]
aps_collection = client["vuln_db"]["aps_scores"]

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

UPLOAD_DIR = "uploads"
TEXT_DIR = "processed_text"
VECTOR_DIR = "vector_store"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEXT_DIR, exist_ok=True)
os.makedirs(VECTOR_DIR, exist_ok=True)

app = FastAPI()

#EMBEDDINGS 
embedder = SentenceTransformer("all-MiniLM-L6-v2")
embedding_dim = 384
index = faiss.IndexFlatL2(embedding_dim)

documents = []   # stores chunk text
doc_ids = []     # stores document IDs

#UTILS
def extract_text(file_path: str, filename: str) -> str:
    if filename.endswith((".png", ".jpg", ".jpeg")):
        text = pytesseract.image_to_string(Image.open(file_path))
        print("OCR RESULT:", text)
        return text

    elif filename.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    elif filename.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        return ""

def chunk_text(text, chunk_size=500, overlap=100):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks

#GOOGLE SEARCH FALLBACK
def google_fallback(query, num_results=5):
    urls = list(search(query + " wikipedia", num_results=num_results))
    contents = []
    for url in urls:
        try:
            r = requests.get(url, timeout=10)
            doc = Document(r.text)
            text = doc.get_clean_html()
            soup = BeautifulSoup(text, "html.parser")
            main_text = soup.get_text(separator="\n")
            contents.append(main_text[:2000])
        except:
            continue
    return "\n\n".join(contents)

#NEW SECTION - APS / PRIORITIZATION MODEL

# Criticality mapping
criticality_map = {"Low": 1, "Medium": 5, "High": 10}

#Example APS formula
def compute_APS(cvss, epss_prob, kev_present, criticality_score):
    return round(
        0.4 * cvss +            # technical severity
        0.3 * epss_prob * 10 +  # real-world exploit probability scaled
        0.2 * (1 if kev_present else 0) +
        0.1 * criticality_score,      # business impact
        2
    )

class VulnScoreRequest(BaseModel):
    cve: str
    cvss: float
    kev: bool
    epss: float
    poc: bool
    criticality: str
    exposure: str
    regulatory_scope: list
    patch_available: bool
    complexity: str
    
@app.post("/score_vuln")
async def score_vuln(req: VulnScoreRequest):
    criticality_score = criticality_map.get(req.criticality, 1)
    aps = compute_APS(req.cvss, req.epss, req.kev, criticality_score)

    aps_collection.insert_one({
        "cve": req.cve,
        "cvss": req.cvss,
        "epss": req.epss,
        "kev": req.kev,
        "criticality": req.criticality,
        "aps": aps,
        "timestamp": datetime.now(timezone.utc)
    })

    return {
        "cve": req.cve,
        "APS_score": aps,
        "explanation_reference": f"/explain/{req.cve}"
    }
    
#NEW SECTION - XAI EXPLANATION ENDPOINT
@app.get("/explain/{cve_id}")
async def explain_score(cve_id: str):

    record = aps_collection.find_one({"cve": cve_id})
    if not record:
        return {"error": "CVE not scored"}

    text = f"""
Explain why the APS score was assigned.

CVSS Score: {record['cvss']}
EPSS Probability: {record['epss']}
KEV Present: {record['kev']}
Criticality: {record['criticality']}
Final APS Score: {record['aps']}
    """

    response = requests.post(
        "http://localhost:11434/api/generate",
        json={"model": "llama3","prompt": text,"stream": False}
    )

    return {
        "cve": cve_id,
        "aps_score": record["aps"],
        "explanation": response.json().get("response","")
    }

#NEW SECTION - METRICS LOGGING

@app.post("/log_metrics")
async def log_metrics(total_vulns:int, total_high:int, mttr:float):

    metrics_collection.insert_one({
        "total_vulnerabilities": total_vulns,
        "high_aps_count": total_high,
        "mean_time_to_remediate": mttr,
        "timestamp": datetime.now(timezone.utc)
    })

    return {"status":"metrics_logged"}

#ORIGINAL RAG ENDPOINTS REMAIN UNCHANGED BELOW

#FILE UPLOAD 
@app.post("/upload")
async def upload_file(file: UploadFile):
    doc_id = str(uuid.uuid4())
    file_path = f"{UPLOAD_DIR}/{doc_id}_{file.filename}"

    with open(file_path, "wb") as f:
        f.write(await file.read())

    text = extract_text(file_path, file.filename.lower())
    
    if not text.strip():
        return {"error": "No text extracted from file"}

    text_file = f"{TEXT_DIR}/{doc_id}.txt"
    with open(text_file, "w", encoding="utf-8") as f:
        f.write(text)

    chunks = chunk_text(text)
    embeddings = embedder.encode(chunks)

    index.add(np.array(embeddings).astype("float32"))
    documents.extend(chunks)
    doc_ids.extend([doc_id] * len(chunks))

    # save metadata to MongoDB
    collection.insert_one({
    "doc_id": doc_id,
    "filename": file.filename,
    "file_path": file_path,
    "text_file": text_file,
    "chunks_created": len(chunks),
    "uploaded_at": datetime.now(timezone.utc)
    })

    return {
        "message": "File processed successfully",
        "doc_id": doc_id,
        "chunks_created": len(chunks)
    }

#QUESTION ANSWERING
@app.post("/ask")
async def ask_question(question: str):

    q_embedding = embedder.encode([question]).astype("float32")
    D, I = index.search(q_embedding, k=3)

    threshold = 10.0
    top_distance = float(D[0][0])

    # CASE 1: documents selected 
    if top_distance < threshold:

        context = "\n\n".join([documents[i] for i in I[0]])

        prompt = f"""
SYSTEM:
Use the context to answer the question.
If answer not found, respond: "Not found in document."

CONTEXT:
{context}

QUESTION:
{question}
"""

        doc_response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3",
                "prompt": prompt,
                "stream": False
            }
        )

        doc_answer = doc_response.json().get("response","")

        not_found_markers = [
            "not found in document",
            "no relevant",
            "cannot answer",
            "no information"
        ]

        # check if doc failed -> fallback
        if any(m in doc_answer.lower() for m in not_found_markers):

            google_context = google_fallback(question)

            prompt2 = f"""
SYSTEM:
Web-augmented answer. Use only WEB CONTEXT.

WEB CONTEXT:
{google_context}

QUESTION:
{question}
"""

            g_res = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": "llama3",
                    "prompt": prompt2,
                    "stream": False
                }
            )

            return {
                "answer": g_res.json().get("response",""),
                "source": "google_fallback_after_doc",
                "distance": top_distance
            }

        # document answered correctly
        return {
            "answer": doc_answer,
            "source": "documents",
            "distance": top_distance
        }


    #CASE 2: direct google fallback 
    google_context = google_fallback(question)

    prompt_g = f"""
SYSTEM:
Web-augmented answer. Use only WEB CONTEXT.

WEB CONTEXT:
{google_context}

QUESTION:
{question}
"""

    google_response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": "llama3",
            "prompt": prompt_g,
            "stream": False
        }
    )

    return {
        "answer": google_response.json().get("response",""),
        "source": "google",
        "distance": top_distance
    }
